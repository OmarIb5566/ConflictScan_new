"""
DOCX Template Filler
====================
Fills a Word (.docx) template by replacing {{token}} placeholders with values.

Why this exists
---------------
The RFI form is the client's document, not ours. Rather than generating a
layout, we take whatever .docx the project uses, and only substitute the
placeholder tokens - every logo, table, border, and footer in the original
form survives untouched.

Placeholder syntax
------------------
    {{token_name}}

Word frequently splits a single visible string across several XML runs (it does
this whenever spell-check, formatting, or editing history intervenes), so a
naive run-by-run replace misses most placeholders. This module joins the runs
of each paragraph, substitutes on the joined text, then writes the result back
into the first run and blanks the rest - which preserves the paragraph's
formatting while guaranteeing the token is found.

Headers, footers, tables (including nested tables) and text boxes are all
covered.
"""

import os
import re
from typing import Dict, Any, List

from docx.oxml.ns import qn

TOKEN_PATTERN = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")

# The company RFI form's Discipline row uses real Word checkbox content
# controls (w:sdt / w14:checkbox), not text tokens - they can't be filled by
# _substitute(). This ticks the one matching the drafted discipline and
# leaves the rest unchecked, in the order they appear in the form.
DISCIPLINE_CHECKBOX_ORDER = [
    "survey", "civil", "structural", "electrical",
    "mechanical", "plumbing", "arch", "other",
]
DISCIPLINE_ALIASES = {
    "architectural": "arch",
    "architecture": "arch",
    "contractual": "other",
    "general": "other",
}
CHECKBOX_CHECKED = "☒"    # ☒
CHECKBOX_UNCHECKED = "☐"  # ☐


def _normalise_discipline(value: str) -> str:
    key = (value or "").strip().lower()
    key = DISCIPLINE_ALIASES.get(key, key)
    return key if key in DISCIPLINE_CHECKBOX_ORDER else "other"


def _tick_discipline_checkboxes(document, discipline_value: str) -> None:
    """Tick the Discipline checkbox matching *discipline_value*; uncheck the rest."""
    if not discipline_value:
        return
    target = _normalise_discipline(discipline_value)

    checkbox_sdts = [
        sdt for sdt in document.element.body.iter(qn("w:sdt"))
        if sdt.find(".//" + qn("w14:checkbox")) is not None
    ]
    for name, sdt in zip(DISCIPLINE_CHECKBOX_ORDER, checkbox_sdts):
        is_target = name == target
        checkbox = sdt.find(".//" + qn("w14:checkbox"))
        checked_el = checkbox.find(qn("w14:checked")) if checkbox is not None else None
        if checked_el is not None:
            checked_el.set(qn("w14:val"), "1" if is_target else "0")
        glyph_run = sdt.find(".//" + qn("w:t"))
        if glyph_run is not None:
            glyph_run.text = CHECKBOX_CHECKED if is_target else CHECKBOX_UNCHECKED


def _substitute(text: str, values: Dict[str, Any]) -> str:
    """Replace every {{token}} in *text*. Unknown tokens are emptied, not left raw."""
    def repl(match):
        key = match.group(1)
        value = values.get(key, "")
        return "" if value is None else str(value)
    return TOKEN_PATTERN.sub(repl, text)


def _fill_paragraph(paragraph, values: Dict[str, Any]) -> None:
    """Substitute tokens in one paragraph, preserving the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        return

    joined = "".join(run.text for run in runs)
    if "{{" not in joined:
        return

    replaced = _substitute(joined, values)
    if replaced == joined:
        return

    # Multi-line values become real line breaks rather than literal "\n".
    # add_break() and add_t() append XML children in call order, so the run
    # ends up as: text, <br/>, text, <br/>, text ...  Assigning run.text again
    # would wipe those breaks, hence the low-level add_t().
    lines = replaced.split("\n")
    runs[0].text = lines[0]
    for extra in lines[1:]:
        runs[0].add_break()
        runs[0]._r.add_t(extra)

    for run in runs[1:]:
        run.text = ""


def _fill_table(table, values: Dict[str, Any]) -> None:
    """Recursively substitute tokens in every cell of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _fill_paragraph(paragraph, values)
            for nested in cell.tables:
                _fill_table(nested, values)


def _fill_container(container, values: Dict[str, Any]) -> None:
    """Substitute tokens in every paragraph and table of a document or section part."""
    for paragraph in getattr(container, "paragraphs", []):
        _fill_paragraph(paragraph, values)
    for table in getattr(container, "tables", []):
        _fill_table(table, values)


def fill_docx_template(template_path: str, output_path: str, values: Dict[str, Any]) -> str:
    """
    Copy *template_path* to *output_path* with every {{token}} replaced.

    Parameters
    ----------
    template_path : path to the .docx form to fill
    output_path   : where the filled copy is written
    values        : {token_name: value}; tokens absent from this dict are
                    replaced with an empty string so no placeholder text ever
                    reaches the client

    Returns the output path.

    Raises FileNotFoundError if the template is missing, and ImportError if
    python-docx is not installed.
    """
    from docx import Document   # imported here so the module loads without python-docx

    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"RFI template not found: {template_path}")

    document = Document(template_path)

    # Body
    _fill_container(document, values)
    _tick_discipline_checkboxes(document, values.get("discipline"))

    # Headers and footers, per section
    for section in document.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is not None:
                _fill_container(part, values)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    document.save(output_path)
    return output_path


def list_template_tokens(template_path: str) -> List[str]:
    """
    Return every distinct {{token}} name found in a template, in order of first
    appearance. Useful for checking a newly supplied form against the tokens
    the RFI drafting node actually produces.
    """
    from docx import Document

    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"RFI template not found: {template_path}")

    document = Document(template_path)
    found: List[str] = []

    def scan(text: str) -> None:
        for match in TOKEN_PATTERN.finditer(text or ""):
            name = match.group(1)
            if name not in found:
                found.append(name)

    def scan_container(container) -> None:
        for paragraph in getattr(container, "paragraphs", []):
            scan("".join(r.text for r in paragraph.runs))
        for table in getattr(container, "tables", []):
            scan_table(table)

    def scan_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    scan("".join(r.text for r in paragraph.runs))
                for nested in cell.tables:
                    scan_table(nested)

    scan_container(document)
    for section in document.sections:
        for part in (section.header, section.footer):
            if part is not None:
                scan_container(part)

    return found
