"""
Generates the placeholder RFI form shipped with the project.

Run this only to regenerate rfi_template.docx from scratch:

    python templates_rfi/make_sample_template.py

In normal use you replace rfi_template.docx with the project's own RFI form -
keep the {{token}} placeholders, and the form's layout, logo, and footer are
preserved exactly as they are.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


HEADER_FIELDS = [
    ("RFI No.", "{{rfi_no}}"),
    ("Date", "{{rfi_date}}"),
    ("Project", "{{project_name}}"),
    ("Discipline", "{{discipline}}"),
    ("Priority", "{{priority}}"),
    ("Response required by", "{{response_required_by}}"),
    ("Originator", "{{originator}}"),
]

REFERENCE_FIELDS = [
    ("Document A", "{{document_a_label}}"),
    ("Reference", "{{document_a_reference}}"),
    ("Document B", "{{document_b_label}}"),
    ("Reference", "{{document_b_reference}}"),
]

BODY_SECTIONS = [
    ("Subject", "{{subject}}"),
    ("Background", "{{background}}"),
    ("Discrepancy", "{{discrepancy_summary}}"),
    ("Information requested", "{{question}}"),
    ("Proposed way forward", "{{proposed_solution}}"),
    ("Cost impact", "{{cost_impact}}"),
    ("Programme impact", "{{schedule_impact}}"),
]


def build(path: str) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph("REQUEST FOR INFORMATION")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    subtitle = doc.add_paragraph("Ref: {{discrepancy_id}}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Header block -----------------------------------------------------
    header = doc.add_table(rows=len(HEADER_FIELDS), cols=2)
    header.style = "Table Grid"
    for row, (label, token) in zip(header.rows, HEADER_FIELDS):
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = token
    header.columns[0].width = Cm(5)

    doc.add_paragraph()

    # --- Documents in question -------------------------------------------
    heading = doc.add_paragraph("DOCUMENTS IN QUESTION")
    heading.runs[0].bold = True

    refs = doc.add_table(rows=len(REFERENCE_FIELDS), cols=2)
    refs.style = "Table Grid"
    for row, (label, token) in zip(refs.rows, REFERENCE_FIELDS):
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = token
    refs.columns[0].width = Cm(5)

    doc.add_paragraph()

    # --- Body sections ----------------------------------------------------
    for label, token in BODY_SECTIONS:
        p = doc.add_paragraph(label.upper())
        p.runs[0].bold = True
        doc.add_paragraph(token)
        doc.add_paragraph()

    # --- Client response block -------------------------------------------
    response_heading = doc.add_paragraph("CLIENT / ENGINEER RESPONSE")
    response_heading.runs[0].bold = True

    response = doc.add_table(rows=2, cols=2)
    response.style = "Table Grid"
    response.rows[0].cells[0].text = "Response"
    response.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    response.rows[0].cells[1].text = "\n\n\n"
    response.rows[1].cells[0].text = "Signed / Date"
    response.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    response.rows[1].cells[1].text = ""
    response.columns[0].width = Cm(5)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "{{project_name}}  |  {{rfi_no}}  |  Issued {{rfi_date}}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)
    return path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rfi_template.docx")
    print(f"Wrote {build(out)}")
